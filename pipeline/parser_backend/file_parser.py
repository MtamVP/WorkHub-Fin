import io
import traceback

try:
    import pdfplumber
    import docx
    import pandas as pd
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    print("ALL IMPORTS SUCCESSFUL")
except Exception as e:
    print("IMPORT ERROR:", e)
    traceback.print_exc()

def chunk_text(text: str) -> list[str]:
    # Sử dụng Langchain để cắt văn bản thành các chunk nhỏ
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=100,
        length_function=len,
        is_separator_regex=False,
    )
    return text_splitter.split_text(text)

def parse_csv(file_bytes: bytes) -> str:
    # Đọc csv, bỏ qua lỗi encoding hoặc format nếu có thể
    df = pd.read_csv(io.BytesIO(file_bytes), on_bad_lines='skip')
    # Xóa dòng trắng
    df = df.dropna(how='all')
    
    rows = []
    headers = df.columns.tolist()
    
    for idx, row in df.iterrows():
        row_items = []
        for col_name in headers:
            val = row[col_name]
            if pd.notna(val) and str(val).strip() != "":
                row_items.append(f"{col_name}: {val}")
        if row_items:
            rows.append(f"[Dòng {idx + 1}] " + " | ".join(row_items))
    return "\n".join(rows)

def parse_excel(file_bytes: bytes) -> str:
    # Đọc tất cả các sheet
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    all_text = []
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        df = df.dropna(how='all')
        
        all_text.append(f"\n--- Bảng: {sheet_name} ---")
        headers = df.columns.tolist()
        for idx, row in df.iterrows():
            row_items = []
            for col_name in headers:
                val = row[col_name]
                if pd.notna(val) and str(val).strip() != "":
                    # Nếu column name chứa Unnamed thì bỏ qua tiêu đề cột đó (bảng không có header chuẩn)
                    if "Unnamed:" in str(col_name):
                        row_items.append(f"{val}")
                    else:
                        row_items.append(f"{col_name}: {val}")
            if row_items:
                all_text.append(f"[Dòng {idx + 1}] " + " | ".join(row_items))
    return "\n".join(all_text)

def parse_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
            
    # Bóc tách các bảng trong Word (nếu có)
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_data:
                full_text.append(f"[Bảng - Dòng {i+1}] " + " | ".join(row_data))
                
    return "\n\n".join(full_text)

def parse_pdf(file_bytes: bytes) -> str:
    # Sử dụng pdfplumber để trích xuất chữ và bảng biểu chất lượng cao
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            # Lấy chữ thuần
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
            
            # Cố gắng lấy bảng (Table)
            tables = page.extract_tables()
            for table in tables:
                for row_idx, row in enumerate(table):
                    row_data = [str(cell).strip() for cell in row if cell and str(cell).strip()]
                    if row_data:
                        text_parts.append(f"[Trang {i+1} - Bảng - Dòng {row_idx+1}] " + " | ".join(row_data))
                        
    return "\n\n".join(text_parts)

def extract_text(file_bytes: bytes, file_name: str) -> str:
    ext = file_name.split('.')[-1].lower() if '.' in file_name else ''
    
    if ext in ['csv', 'tsv']:
        return parse_csv(file_bytes)
    elif ext in ['xlsx', 'xls']:
        return parse_excel(file_bytes)
    elif ext in ['docx', 'doc']:
        return parse_docx(file_bytes)
    elif ext == 'pdf':
        return parse_pdf(file_bytes)
    else:
        # Xử lý text thô làm mặc định
        return file_bytes.decode('utf-8', errors='ignore')
